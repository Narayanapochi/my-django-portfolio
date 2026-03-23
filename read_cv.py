import docx

def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in full_text:
                    full_text.append(text)
    return '\n'.join(full_text)

print(get_text(r"c:\Users\hp\Downloads\portfolio-master\portfolio-master\Narayana.docx"))
